"""Functions to generate Excel, Word, and Markdown outputs."""

from __future__ import annotations

import json
import logging
import os
import shutil
import tempfile
from datetime import datetime
from pathlib import Path

import requests
import xlsxwriter  # type: ignore[import]
from docx import Document  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]
from docx.shared import Inches, Pt  # type: ignore[import]
from PIL import Image as PILImage  # type: ignore[import]

from modules.utils.blob_utils import upload_content_to_blob, upload_file_to_blob

logger = logging.getLogger(__name__)


def create_excel_with_layout(excel_file_path: str, image_paths: list[str], responses: list[str], new_height: int = 400) -> str:
    """
    画像と対応するテキスト応答を含むExcelファイルを作成する関数

    Args:
        excel_file_path (str): 作成するExcelファイルのパス
        image_paths (list): Excelに含める画像ファイルのパスのリスト
        responses (list): 各画像に対応するテキスト応答のリスト
        new_height (int): 画像の新しい高さ（ピクセル単位）。デフォルトは400ピクセル。

    Returns:
        str: 生成されたExcelファイルのパス

    Raises:
        ValueError: 画像パスとテキストのリストの長さが一致しない場合に発生
    """
    if len(image_paths) != len(responses):
        raise ValueError("画像パスとテキストのリストの長さが一致していません。")

    output_dir = Path(excel_file_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    workbook = xlsxwriter.Workbook(excel_file_path)
    worksheet = workbook.add_worksheet()
    text_format = workbook.add_format({"text_wrap": True, "valign": "top", "font_name": "Yu Gothic UI", "font_size": 12})
    row = 0

    # 画像を指定された高さにリサイズし、Excelに挿入
    for i, (image_path, response) in enumerate(zip(image_paths, responses)):
        try:
            with PILImage.open(image_path) as img:
                h_percent = new_height / float(img.size[1])
                new_width = int(float(img.size[0]) * h_percent)

                resample_filter = PILImage.LANCZOS  # type: ignore[attr-defined]
                img = img.resize((new_width, new_height), resample=resample_filter)
                img_width, img_height = img.size
                temp_img_path = f"{str(output_dir)}/temp_resized_image_{i}.png"
                img.save(temp_img_path, format="PNG")

            worksheet.set_row_pixels(row=row, height=img_height + 20)
            if i == 0:
                worksheet.set_column_pixels(first_col=0, last_col=1, width=img_width + 20)

            worksheet.insert_image(row, 0, temp_img_path, {"x_offset": 5, "y_offset": 5})
        except Exception:
            continue
        worksheet.write(row, 1, response, text_format)
        row += 1
    workbook.close()

    # 一時的なリサイズ画像を削除
    for i in range(len(image_paths)):
        tmp = Path(f"{str(output_dir)}/temp_resized_image_{i}.png")
        if tmp.exists():
            tmp.unlink()
    return excel_file_path


def create_markdown_with_table(md_file_path: str, image_paths: list[str], responses: list[str]) -> str:
    """
    画像と対応するテキストを含むMarkdownファイルを作成する関数

    Args:
        md_file_path (str): 作成するMarkdownファイルのパス
        image_paths (list): Markdownに含める画像ファイルのパスのリスト
        responses (list): 各画像に対応するテキストのリスト

    Returns:
        str: 生成されたMarkdownファイルのパス。

    Raises:
        ValueError: 画像パスとテキストのリストの長さが一致しない場合に発生
    """
    if len(image_paths) != len(responses):
        raise ValueError("画像パスとテキストのリストの長さが一致していません。")

    output_dir = Path(md_file_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    image_dir = Path(output_dir, "images")
    image_dir.mkdir(parents=True, exist_ok=True)

    # 画像をMarkdownディレクトリ内のimagesフォルダにコピー
    img_dest_absolute_paths = [f"{image_dir}/{Path(img_path).name}" for img_path in image_paths]
    for img_path, img_dest in zip(image_paths, img_dest_absolute_paths):
        shutil.copy(img_path, img_dest)

    # 画像の相対パスを作成
    img_dest_relative_paths = [f"./images/{Path(img_path).name}" for img_path in image_paths]

    # Markdownテーブルを作成
    markdown_lines = ["| 画像 | 説明 |", "|------|------|"]
    for img_path, response_text in zip(img_dest_relative_paths, responses):
        replaced_text = response_text.replace("\n", "<br />")
        markdown_lines.append(f"| ![]({img_path}) | {replaced_text} |")

    # Markdownファイルを書き込み
    content = "\n".join(markdown_lines)
    with open(md_file_path, "w", encoding="utf-8") as f:
        f.write(content)
    return md_file_path


def create_word_with_layout(word_file_path: str, image_paths: list[str], responses: list[str], max_width: float = 6) -> str:
    """
    画像と対応するテキスト応答を含むWordファイルを作成する関数

    Args:
        word_file_path (str): 作成するWordファイルのパス
        image_paths (list): Wordに含める画像ファイルのパスのリスト
        responses (list): 各画像に対応するテキスト応答のリスト
        max_width (float): 画像の最大幅（インチ単位）。デフォルトは6インチ。

    Returns:
        str: 生成されたWordファイルのパス

    Raises:
        ValueError: 画像パスとテキストのリストの長さが一致しない場合に発生
    """
    if len(image_paths) != len(responses):
        raise ValueError("画像パスとテキストのリストの長さが一致していません。")

    output_dir = Path(word_file_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    # Wordドキュメントを作成し、画像とテキストを追加
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "MS Mincho"
    font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "MS Mincho")  # type: ignore[attr-defined]
    for i, (image_path, response) in enumerate(zip(image_paths, responses)):
        table = doc.add_table(rows=2, cols=1)
        table.autofit = False
        try:
            with PILImage.open(image_path) as img:
                aspect_ratio = img.height / img.width
                max_height = max_width * aspect_ratio
                resized_path = f"{str(output_dir)}/temp_resized_image_{i}.png"
                img = img.resize((int(96 * max_width), int(96 * max_height)))
                img.save(resized_path)
                paragraph = table.rows[0].cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(resized_path, width=Inches(max_width))
        except Exception:
            continue
        table.rows[1].cells[0].text = response
        doc.add_page_break()

    # Wordファイルを保存
    doc.save(word_file_path)
    for i in range(len(image_paths)):
        tmp = Path(f"{str(output_dir)}/temp_resized_image_{i}.png")
        if tmp.exists():
            tmp.unlink()
    return word_file_path


def update_excel(steps, frame_urls, container_name, blob_folder_name, correlation_id):
    """Excelファイルを更新"""
    try:
        logger.info("📊 Excelファイル更新開始")
        logger.info(
            f"📋 更新対象データ: steps={len(steps)}件, frame_urls={len(frame_urls)}件"
        )

        # ステップデータの詳細ログ
        for idx, step in enumerate(steps[:5]):
            logger.info(
                f"  📝 Step {idx+1}: id={step.id}, frameIdx={step.frameIdx}, description='{step.description[:50]}...'"
            )
        if len(steps) > 5:
            logger.info(f"  📝 ... and {len(steps) - 5} more steps")

        # 一時ディレクトリを作成
        temp_dir = tempfile.mkdtemp()

        # 画像パスとレスポンスを準備
        response_image_paths = []
        responses = []

        for step in steps:
            frame_idx = step.frameIdx
            if frame_idx < len(frame_urls):
                # 画像URLをローカルファイルとしてダウンロード
                frame_url = frame_urls[frame_idx]
                logger.info(
                    f"  🖼️ Processing step {step.id}: frameIdx={frame_idx}, URL={frame_url[:100]}..."
                )
                try:
                    # URLから一時ファイルにダウンロード
                    response = requests.get(frame_url)
                    if response.status_code == 200:
                        temp_image_path = Path(temp_dir) / f"frame_{frame_idx}.jpg"
                        with open(temp_image_path, "wb") as f:
                            f.write(response.content)
                        response_image_paths.append(str(temp_image_path))
                        responses.append(step.description)
                        logger.info(f"  ✅ Step {step.id}: 画像ダウンロード成功")
                except Exception as e:
                    logger.warning(
                        f"⚠️ フレーム画像ダウンロード失敗: {frame_url}, エラー: {str(e)}"
                    )
            else:
                logger.warning(
                    f"⚠️ Step {step.id}: frameIdx {frame_idx} が frame_urls の範囲外 (max: {len(frame_urls)-1})"
                )

        logger.info(
            f"📊 Excel生成用データ準備完了: images={len(response_image_paths)}, responses={len(responses)}"
        )

        # Excelファイル生成（タイムスタンプ付きファイル名）
        file_prefix = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
        excel_file_blob_name = f"{blob_folder_name}/excels/{file_prefix}-guideline.xlsx"
        excel_local_file_path = Path(temp_dir) / "excel_output.xlsx"
        excel_file_path = create_excel_with_layout(
            excel_file_path=str(excel_local_file_path),
            image_paths=response_image_paths,
            responses=responses,
            new_height=400,
        )

        # Blobにアップロード
        excel_url = upload_file_to_blob(
            excel_file_path, container_name, excel_file_blob_name
        )

        # 一時ファイルクリーンアップ
        shutil.rmtree(temp_dir)

        logger.info("✅ Excelファイル更新完了")
        return excel_url

    except Exception as e:
        logger.error(f"❌ Excelファイル更新エラー: {str(e)}")
        if "temp_dir" in locals():
            shutil.rmtree(temp_dir)
        return None


def update_docx(steps, frame_urls, container_name, blob_folder_name, correlation_id):
    """Wordファイルを更新"""
    try:
        logger.info("📝 Wordファイル更新開始")

        # 一時ディレクトリを作成
        temp_dir = tempfile.mkdtemp()

        # 画像パスとレスポンスを準備
        response_image_paths = []
        responses = []

        for step in steps:
            frame_idx = step.frameIdx
            if frame_idx < len(frame_urls):
                frame_url = frame_urls[frame_idx]
                try:
                    response = requests.get(frame_url)
                    if response.status_code == 200:
                        temp_image_path = Path(temp_dir) / f"frame_{frame_idx}.jpg"
                        with open(temp_image_path, "wb") as f:
                            f.write(response.content)
                        response_image_paths.append(str(temp_image_path))
                        responses.append(step.description)
                except Exception as e:
                    logger.warning(
                        f"⚠️ フレーム画像ダウンロード失敗: {frame_url}, エラー: {str(e)}"
                    )

        # Wordファイル生成（タイムスタンプ付きファイル名）
        file_prefix = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
        word_file_blob_name = f"{blob_folder_name}/words/{file_prefix}-guideline.docx"
        word_local_file_path = Path(temp_dir) / "word_output.docx"
        word_file_path = create_word_with_layout(
            word_file_path=str(word_local_file_path),
            image_paths=response_image_paths,
            responses=responses,
            max_width=6,
        )

        # Blobにアップロード
        word_url = upload_file_to_blob(
            word_file_path, container_name, word_file_blob_name
        )

        # 一時ファイルクリーンアップ
        shutil.rmtree(temp_dir)

        logger.info("✅ Wordファイル更新完了")
        return word_url

    except Exception as e:
        logger.error(f"❌ Wordファイル更新エラー: {str(e)}")
        if "temp_dir" in locals():
            shutil.rmtree(temp_dir)
        return None


def update_markdown(steps, frame_urls, container_name, blob_folder_name, correlation_id):
    """Markdownファイルを更新"""
    try:
        logger.info("📋 Markdownファイル更新開始")

        # 一時ディレクトリを作成
        temp_dir = tempfile.mkdtemp()

        # 画像パスとレスポンスを準備
        response_image_paths = []
        responses = []

        for step in steps:
            frame_idx = step.frameIdx
            if frame_idx < len(frame_urls):
                frame_url = frame_urls[frame_idx]
                try:
                    response = requests.get(frame_url)
                    if response.status_code == 200:
                        temp_image_path = Path(temp_dir) / f"frame_{frame_idx}.jpg"
                        with open(temp_image_path, "wb") as f:
                            f.write(response.content)
                        response_image_paths.append(str(temp_image_path))
                        responses.append(step.description)
                except Exception as e:
                    logger.warning(
                        f"⚠️ フレーム画像ダウンロード失敗: {frame_url}, エラー: {str(e)}"
                    )

        # Markdownファイル生成（タイムスタンプ付きファイル名）
        file_prefix = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
        markdown_file_blob_name = (
            f"{blob_folder_name}/markdowns/{file_prefix}-guideline.md"
        )
        markdown_local_file_path = Path(temp_dir) / "markdown_output.md"
        markdown_file_path = create_markdown_with_table(
            md_file_path=str(markdown_local_file_path),
            image_paths=response_image_paths,
            responses=responses,
        )

        # Blobにアップロード
        markdown_url = upload_file_to_blob(
            markdown_file_path, container_name, markdown_file_blob_name
        )

        # 画像ファイルもアップロード
        file_prefix = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
        for i, image_path in enumerate(response_image_paths):
            try:
                image_file_name = f"frame_{i}.jpg"
                markdown_image_blob_name = f"{blob_folder_name}/markdowns/{file_prefix}-images/{image_file_name}"
                _ = upload_file_to_blob(
                    image_path, container_name, markdown_image_blob_name
                )
            except Exception as e:
                logger.error(f"❌ Markdown画像アップロードエラー: {str(e)}")

        # 一時ファイルクリーンアップ
        shutil.rmtree(temp_dir)

        logger.info("✅ Markdownファイル更新完了")
        return markdown_url

    except Exception as e:
        logger.error(f"❌ Markdownファイル更新エラー: {str(e)}")
        if "temp_dir" in locals():
            shutil.rmtree(temp_dir)
        return None

def update_json(
    steps,
    frame_urls,
    container_name,
    blob_folder_name,
    llm_output_result_blob_name,
    correlation_id,
):
    """JSONファイルを更新"""
    try:
        logger.info("📄 JSONファイル更新開始")

        # ステップデータをJSON形式に変換
        json_records = []
        for step in steps:
            # frameIdxを使用して対応する画像パスを取得
            frame_idx = step.frameIdx
            if frame_idx < len(frame_urls):
                # Blob内のパスに変換
                frame_url = frame_urls[frame_idx]
                storage_account = os.environ.get("AZURE_STORAGE_ACCOUNT_NAME", "")
                blob_img_path = frame_url.replace(
                    f"https://{storage_account}.blob.core.windows.net/{container_name}/",
                    "",
                )
                json_records.append(
                    {
                        "image": blob_img_path,
                        "text": step.description,
                    }
                )

        # JSONファイルをBlobにアップロード
        json_content = json.dumps(json_records, ensure_ascii=False, indent=2)
        upload_content_to_blob(
            json_content, container_name, llm_output_result_blob_name
        )

        logger.info("✅ JSONファイル更新完了")
        return True

    except Exception as e:
        logger.error(f"❌ JSONファイル更新エラー: {str(e)}")
        return False

